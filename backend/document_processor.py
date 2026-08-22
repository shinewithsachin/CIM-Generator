"""
Universal document processor.
Supports: PDF, DOCX, XLSX, CSV, XML, TXT, JSON, images (OCR), URLs
"""
import json
import pandas as pd
from pathlib import Path
from typing import List
from langchain.schema import Document


class DocumentProcessor:
    """Parse any file format into a list of LangChain Documents."""

    SUPPORTED_EXTENSIONS = {
        ".pdf", ".docx", ".doc",
        ".xlsx", ".xls", ".csv",
        ".xml", ".json", ".txt",
        ".png", ".jpg", ".jpeg", ".tiff", ".bmp", ".gif",
        ".html", ".htm",
    }

    def process_file(self, file_path: str) -> List[Document]:
        ext = Path(file_path).suffix.lower()
        name = Path(file_path).name

        if ext == ".pdf":
            return self._pdf(file_path, name)
        elif ext in (".docx", ".doc"):
            return self._docx(file_path, name)
        elif ext in (".xlsx", ".xls"):
            return self._xlsx(file_path, name)
        elif ext == ".csv":
            return self._csv(file_path, name)
        elif ext == ".xml":
            return self._xml(file_path, name)
        elif ext == ".json":
            return self._json(file_path, name)
        elif ext == ".txt":
            return self._txt(file_path, name)
        elif ext in (".html", ".htm"):
            return self._html(file_path, name)
        elif ext in (".png", ".jpg", ".jpeg", ".tiff", ".bmp", ".gif"):
            return self._image(file_path, name)
        else:
            # Try as plain text fallback
            return self._txt(file_path, name)

    def process_url(self, url: str) -> List[Document]:
        try:
            import trafilatura
            downloaded = trafilatura.fetch_url(url)
            text = trafilatura.extract(downloaded, include_tables=True) or ""
            return self._split_text(text, {"source": url, "type": "url"})
        except Exception as e:
            return [Document(page_content=f"Failed to fetch URL {url}: {e}", metadata={"source": url})]

    # ─────────────────────────────────────────────
    # Format-specific parsers
    # ─────────────────────────────────────────────

    def _pdf(self, path: str, name: str) -> List[Document]:
        docs = []
        try:
            from pypdf import PdfReader
            reader = PdfReader(path)
            for i, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                if text.strip():
                    docs.append(Document(
                        page_content=text,
                        metadata={"source": name, "page": i + 1, "type": "pdf"}
                    ))
        except Exception as e:
            docs.append(Document(page_content=f"Error reading PDF {name}: {e}", metadata={"source": name}))
        return docs

    def _docx(self, path: str, name: str) -> List[Document]:
        docs = []
        try:
            import docx as python_docx
            doc = python_docx.Document(path)
            full_text = []

            # Paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)

            # Tables
            for table in doc.tables:
                rows = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    rows.append(" | ".join(row_data))
                table_text = "\n".join(rows)
                if table_text.strip():
                    full_text.append(f"\n[TABLE]\n{table_text}\n[/TABLE]")

            text = "\n".join(full_text)
            docs = self._split_text(text, {"source": name, "type": "docx"})
        except Exception as e:
            docs.append(Document(page_content=f"Error reading DOCX {name}: {e}", metadata={"source": name}))
        return docs

    def _xlsx(self, path: str, name: str) -> List[Document]:
        docs = []
        try:
            xl = pd.ExcelFile(path)
            for sheet_name in xl.sheet_names:
                df = xl.parse(sheet_name)
                # Convert to readable text
                text_rows = [f"Sheet: {sheet_name}"]
                text_rows.append(df.to_string(index=False, max_rows=500))

                # Also try to extract numeric summaries
                numeric_cols = df.select_dtypes(include='number').columns.tolist()
                if numeric_cols:
                    text_rows.append(f"\nNumeric Summary for sheet '{sheet_name}':")
                    text_rows.append(df[numeric_cols].describe().to_string())

                text = "\n".join(text_rows)
                docs.extend(self._split_text(text, {"source": name, "sheet": sheet_name, "type": "xlsx"}))
        except Exception as e:
            docs.append(Document(page_content=f"Error reading XLSX {name}: {e}", metadata={"source": name}))
        return docs

    def _csv(self, path: str, name: str) -> List[Document]:
        try:
            df = pd.read_csv(path)
            text = df.to_string(index=False, max_rows=1000)
            numeric_cols = df.select_dtypes(include='number').columns.tolist()
            if numeric_cols:
                text += f"\n\nNumeric Summary:\n{df[numeric_cols].describe().to_string()}"
            return self._split_text(text, {"source": name, "type": "csv"})
        except Exception as e:
            return [Document(page_content=f"Error reading CSV {name}: {e}", metadata={"source": name})]

    def _xml(self, path: str, name: str) -> List[Document]:
        try:
            import xmltodict
            with open(path, "r", encoding="utf-8", errors="replace") as f:
                content = f.read()
            parsed = xmltodict.parse(content)
            text = json.dumps(parsed, indent=2)
            return self._split_text(text, {"source": name, "type": "xml"})
        except Exception as e:
            return [Document(page_content=f"Error reading XML {name}: {e}", metadata={"source": name})]

    def _json(self, path: str, name: str) -> List[Document]:
        try:
            with open(path, "r", encoding="utf-8", errors="replace") as f:
                data = json.load(f)
            text = json.dumps(data, indent=2)
            return self._split_text(text, {"source": name, "type": "json"})
        except Exception as e:
            return [Document(page_content=f"Error reading JSON {name}: {e}", metadata={"source": name})]

    def _txt(self, path: str, name: str) -> List[Document]:
        try:
            with open(path, "r", encoding="utf-8", errors="replace") as f:
                text = f.read()
            return self._split_text(text, {"source": name, "type": "txt"})
        except Exception as e:
            return [Document(page_content=f"Error reading TXT {name}: {e}", metadata={"source": name})]

    def _html(self, path: str, name: str) -> List[Document]:
        try:
            from lxml import html as lxml_html
            with open(path, "r", encoding="utf-8", errors="replace") as f:
                content = f.read()
            tree = lxml_html.fromstring(content)
            text = tree.text_content()
            return self._split_text(text, {"source": name, "type": "html"})
        except Exception as e:
            return [Document(page_content=f"Error reading HTML {name}: {e}", metadata={"source": name})]

    def _image(self, path: str, name: str) -> List[Document]:
        try:
            import pytesseract
            from PIL import Image
            img = Image.open(path)
            text = pytesseract.image_to_string(img)
            if text.strip():
                return self._split_text(text, {"source": name, "type": "image_ocr"})
            return [Document(page_content=f"[Image: {name} — no text extracted via OCR]", metadata={"source": name})]
        except ImportError:
            return [Document(page_content=f"[Image: {name} — pytesseract not available for OCR]", metadata={"source": name})]
        except Exception as e:
            return [Document(page_content=f"Error processing image {name}: {e}", metadata={"source": name})]

    # ─────────────────────────────────────────────
    # Text chunking
    # ─────────────────────────────────────────────

    def _split_text(self, text: str, metadata: dict, chunk_size: int = 1000, overlap: int = 200) -> List[Document]:
        from langchain.text_splitter import RecursiveCharacterTextSplitter
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=overlap,
            separators=["\n\n", "\n", ". ", " ", ""],
        )
        chunks = splitter.split_text(text)
        return [
            Document(page_content=chunk, metadata={**metadata, "chunk": i})
            for i, chunk in enumerate(chunks) if chunk.strip()
        ]
